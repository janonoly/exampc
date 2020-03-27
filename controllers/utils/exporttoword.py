import docx
import os
from PIL import Image
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm, Inches
from sqlalchemy import and_
from sqlalchemy.orm import sessionmaker
from model.createdb import engine
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class ExportToWord(object):
    def __init__(self,filepath,filename,coursename):
        self.filepath=filepath
        self.filename = filename
        self.coursename = coursename
    def exportpaper(self,papername,questionidlist):
        self.file = docx.Document()  # 创建内存中的word文档对象
        self.file.styles['Normal'].font.name = u'宋体'
        # file.add_paragraph("窗前明月光")  # 写入若干段落
        papertitle =    self.coursename  + '卷'+papername

        self.titlestyle( papertitle)
        # self.erjistyle('不知道')
        # self.contentstyle('内容')
        tihao=0

        paperlistset=self.getpaperset()
        xznum=paperlistset.single_choice_num
        xzfenshu=paperlistset.single_choice_score
        pdnum=paperlistset.judgment
        pdfenshu=paperlistset.judgment_score
        mxznum=paperlistset.multiple_choice_num
        mxzfenshu=paperlistset.multiple_choice_score
        tknum=paperlistset.tk_choice_num
        tkfenshu=paperlistset.tk_choice_score
        jdnum = paperlistset.jd_choice_num
        jdfenshu = paperlistset.jd_choice_score
        mcjsnum = paperlistset.mcjs_choice_num
        mcjsfenshu = paperlistset.mcjs_choice_score
        shijian=paperlistset.kaoshishijian
        self.shijianstyle('(考试时间：'+str(shijian)+'分钟)')
        if xznum:
            self.erjistyle('选择题（每题%s分）'%xzfenshu)
            for i in range(xznum):
                tihao+=1
                id=questionidlist[i].id

                self.genwordxzmxzstr(tihao,id)

        if pdnum:
            self.erjistyle('判断题（每题%s分）' % pdfenshu)
            for i in range(xznum,xznum+pdnum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordpdjdstr(tihao,id)

        if mxznum:
            self.erjistyle('多选题（每题%s分）' % mxzfenshu)
            for i in range(xznum+pdnum,xznum+pdnum+mxznum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordxzmxzstr(tihao,id)

        if tknum:
            self.erjistyle('填空题（每题%s分）' % tkfenshu)
            for i in range(xznum+pdnum+mxznum, xznum+pdnum+mxznum+tknum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordpdjdstr(tihao,id)
        if mcjsnum:
            self.erjistyle('填空题（每题%s分）' % mcjsfenshu)
            for i in range(xznum + pdnum + mxznum + tknum, xznum + pdnum + mxznum + tknum + mcjsnum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordpdjdstr(tihao, id)
        if jdnum:
            self.erjistyle('简答题（每题%s分）' % jdfenshu)
            for i in range(xznum + pdnum + mxznum+tknum+mcjsnum, xznum + pdnum + mxznum + tknum+mcjsnum+jdnum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordpdjdstr(tihao, id)


        filename = self.filepath + '/' +  papertitle+self.filename
        self.file.save(filename)  # 保存才能看到结果
        self.exportpaperanswer(papertitle,questionidlist)

    def exportpaperanswer(self, papertitle, questionidlist):
        self.file = docx.Document()  # 创建内存中的word文档对象

        self.file.styles['Normal'].font.name = u'宋体'
        # file.add_paragraph("窗前明月光")  # 写入若干段落

        self.titlestyle(papertitle + '答案')
        # self.erjistyle('不知道')
        # self.contentstyle('内容')
        tihao = 0

        paperlistset = self.getpaperset()
        xznum = paperlistset.single_choice_num
        xzfenshu = paperlistset.single_choice_score
        pdnum = paperlistset.judgment
        pdfenshu = paperlistset.judgment_score
        mxznum = paperlistset.multiple_choice_num
        mxzfenshu = paperlistset.multiple_choice_score
        jdnum = paperlistset.jd_choice_num
        jdfenshu = paperlistset.jd_choice_score
        mcjsnum = paperlistset.mcjs_choice_num
        mcjsfenshu = paperlistset.mcjs_choice_score

        if xznum:
            xzstr = ''
            self.erjistyle('选择题（每题%s分）' % xzfenshu)
            for i in range(xznum):
                tihao += 1
                id = questionidlist[i].id
                single_question_set = self.gensinglequestion(id)
                xzstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(xzstr)
        if pdnum:
            pdstr = ''
            self.erjistyle('判断题（每题%s分）' % pdfenshu)
            for i in range(xznum, xznum + pdnum):
                tihao += 1
                id = questionidlist[i].id
                single_question_set = self.gensinglequestion(id)
                pdstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(pdstr)

        if mxznum:
            mxzstr = ''
            self.erjistyle('多选题（每题%s分）' % mxzfenshu)
            for i in range(xznum + pdnum, xznum + pdnum + mxznum):
                tihao += 1
                id = questionidlist[i].id
                single_question_set = self.gensinglequestion(id)
                mxzstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(mxzstr)
        if mcjsnum:
            mcjsstr = ''
            self.erjistyle('名词解释（每题%s分）' % jdfenshu)
            for i in range(xznum + pdnum + mxznum, xznum + pdnum + mxznum + mcjsnum):
                tihao += 1
                id = questionidlist[i].id
                single_question_set = self.gensinglequestion(id)
                mcjsstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(mcjsstr)

        if jdnum:
            jdstr = ''
            self.erjistyle('填空题（每题%s分）' % jdfenshu)
            for i in range(xznum + pdnum + mxznum+mcjsnum, xznum + pdnum + mxznum +mcjsnum+ jdnum):
                tihao += 1
                id = questionidlist[i].id
                single_question_set = self.gensinglequestion(id)
                jdstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(jdstr)


        filename = self.filepath + '/' + papertitle + '答案' + self.filename
        self.file.save(filename)  # 保存才能看到结果

    def titlestyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.bold = True
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size=Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def erjistyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    def contentstyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.first_line_indent = Cm(0.74)
    def shijianstyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def getpaperset(self):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import PaperList
        question_result = None
        try:
            question_result = session.query(PaperList).filter(PaperList.course_name==self.coursename).first()

        except:
            pass

        session.close()
        return question_result

    def gensinglequestion(self,id):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import question
        question_result=None
        try:
            question_result = session.query(question).filter(question.id==id).first()

        except:
            pass

        session.close()
        return question_result
    def genwordpdjdstr(self,tihao,id):
        single_question_set = self.gensinglequestion(id)
        self.contentstyle(str(tihao) + ':' + single_question_set.content)
    def genwordxzmxzstr(self,tihao,id):
        single_question_set = self.gensinglequestion(id)
        self.contentstyle(str(tihao) + ':' + single_question_set.content)
        xuanxiangcontent = ''
        if len(single_question_set.choice_a) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_a
        if len(single_question_set.choice_b) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_b
        if len(single_question_set.choice_c) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_c
        if len(single_question_set.choice_d) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_d
        if len(single_question_set.choice_e) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_e
        if len(single_question_set.choice_f) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_f
        if len(single_question_set.choice_g) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_g
        if len(single_question_set.choice_h) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_h
        self.contentstyle(xuanxiangcontent)

class RandExportToWord(object):
    def __init__(self,filepath,filename , fenzhi):
        self.filepath=filepath
        self.filename = filename
        self.fenzhi = fenzhi
    def exportpaper(self,papername,questionidlist,shijuanname):

        self.file = docx.Document()  # 创建内存中的word文档对象
        self.file.styles['Normal'].font.name = u'宋体'
        # file.add_paragraph("窗前明月光")  # 写入若干段落
        papertitle =shijuanname+ '卷'+papername

        self.titlestyle( papertitle)
        # self.erjistyle('不知道')
        # self.contentstyle('内容')
        tihao=0

        try:
            self.xznum=self.getquestiontypenum(questionidlist,'xz')

            self.pdnum=self.getquestiontypenum(questionidlist,'pd')

            self.mxznum=self.getquestiontypenum(questionidlist,'mxz')

            self.tknum=self.getquestiontypenum(questionidlist,'tk')

            self.jdnum = self.getquestiontypenum(questionidlist,'jd')

            self.mcjsnum = self.getquestiontypenum(questionidlist, 'mcjs')

        except:
            pass
        self.shijianstyle('(考试时间：'+str(self.fenzhi['时间'])+'分钟)')
        dabiaoti=0
        dabiaotidict=['一','二','三','四','五','六']
        if self.xznum>0:
            self.erjistyle('%s、选择题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti],self.fenzhi['单选题'],self.xznum))
            dabiaoti += 1
            for i in range(len(questionidlist)):
                id=questionidlist[i].id
                if self.gensinglequestion(id).questionType == 'xz':
                    tihao += 1
                    self.genwordxzmxzstr(tihao,id)

        if self.pdnum > 0:
            self.erjistyle('%s、判断题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['判断题'], self.pdnum))
            dabiaoti += 1
            for i in range(len(questionidlist)):
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType == 'pd':
                    tihao += 1
                    self.genwordxzmxzstr(tihao,id)
        if self.mxznum > 0:
            self.erjistyle('%s、多选题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['多选题'], self.mxznum))
            dabiaoti += 1
            for i in range(len(questionidlist)):

                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'mxz':
                    tihao += 1
                    self.genwordxzmxzstr(tihao,id)
        if self.tknum > 0:
            self.erjistyle('%s、填空题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['填空题'], self.tknum))
            dabiaoti += 1
            for i in range(len(questionidlist)):
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'tk':
                    tihao += 1
                    self.genwordxzmxzstr(tihao,id)

        if self.mcjsnum > 0:
            self.erjistyle('%s、名词解释（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['名词解释'], self.mcjsnum))

            for i in range(len(questionidlist)):
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'mcjs':
                    tihao += 1
                    self.genwordxzmxzstr(tihao, id)

        if self.jdnum > 0:
            self.erjistyle('%s、简答题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['简答题'], self.jdnum))

            for i in range(len(questionidlist)):
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType == 'jd':
                    tihao += 1
                    self.genwordxzmxzstr(tihao, id)

        filename = self.filepath + '/' +  papertitle+self.filename
        self.file.save(filename)  # 保存才能看到结果
        self.exportpaperanswer(papertitle,questionidlist)
    def getquestiontypenum(self,questionidlist,quesitiontype):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import question
        questionnum = 0
        try:
            for i in range(len(questionidlist)):
                questionres = session.query(question).filter(and_(question.id == questionidlist[i].id,question.questionType == quesitiontype)).first()
                if questionres:
                    questionnum+=1

        except:
            pass

        session.close()
        return questionnum
    def exportpaperanswer(self, papertitle, questionidlist):
        self.file = docx.Document()  # 创建内存中的word文档对象

        self.file.styles['Normal'].font.name = u'宋体'
        # file.add_paragraph("窗前明月光")  # 写入若干段落

        self.titlestyle(papertitle + '答案')
        # self.erjistyle('不知道')
        # self.contentstyle('内容')
        tihao = 0


        xzstr = ''
        dabiaoti = 0
        dabiaotidict = ['一', '二', '三', '四', '五', '六']
        if self.xznum > 0:
            self.erjistyle('%s、单选题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['单选题'], self.xznum))
            dabiaoti += 1
            for i in range(len(questionidlist)):
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'xz':
                    tihao += 1
                    single_question_set = self.gensinglequestion(id)
                    xzstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(xzstr)
        pdstr = ''
        if self.pdnum > 0:
            self.erjistyle('%s、判断题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['判断题'], self.pdnum))
            dabiaoti += 1
            for i in range(len(questionidlist)):


                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType == 'pd':
                    tihao += 1

                    single_question_set = self.gensinglequestion(id)
                    pdstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(pdstr)
        mxzstr = ''
        if self.mxznum > 0:
            self.erjistyle('%s、多选题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['多选题'], self.mxznum))
            dabiaoti += 1
            for i in range(len(questionidlist)):


                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'mxz':
                    tihao += 1
                    single_question_set = self.gensinglequestion(id)
                    mxzstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(mxzstr)
        tkstr = ''
        if self.tknum > 0:
            self.erjistyle('%s、填空题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['填空题'], self.tknum))
            dabiaoti += 1
            for i in range(len(questionidlist)):


                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'tk':
                    tihao += 1
                    single_question_set = self.gensinglequestion(id)
                    tkstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(tkstr)

        mcjsstr = ''
        if self.mcjsnum > 0:
            self.erjistyle('%s、名词解释（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['名词解释'], self.mcjsnum))
            for i in range(len(questionidlist)):

                tihao += 1
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType == 'mcjs':
                    single_question_set = self.gensinglequestion(id)
                    mcjsstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(mcjsstr)

        jdstr = ''
        if self.jdnum > 0:
            self.erjistyle('%s、简答题（每题%s分,共%s题）' % (dabiaotidict[dabiaoti], self.fenzhi['简答题'], self.jdnum))
            for i in range(len(questionidlist)):

                tihao += 1
                id = questionidlist[i].id
                if self.gensinglequestion(id).questionType  == 'jd':
                    single_question_set = self.gensinglequestion(id)
                    jdstr += str(tihao) + ':' + single_question_set.answer + '  '
            self.contentstyle(jdstr)

        filename = self.filepath + '/' + papertitle + '答案' + self.filename
        self.file.save(filename)  # 保存才能看到结果

    def titlestyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.bold = True
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size=Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def erjistyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    def contentstyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.first_line_indent = Cm(0.74)
    def contentimgstyle(self,imagespath):
        # p = self.file.add_paragraph()
        # self.file.add_picture(imagespath, width=Inches(8),height=Inches(8))
        img = Image.open(imagespath)
        width,height=img.size
        w=width+height
        dw=width/w*7
        dh=height/w*7
        self.file.add_picture(imagespath, width=Cm(dw), height=Cm(dh))


    def shijianstyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    def gensinglequestion(self,id):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import question
        question_result=None
        try:
            question_result = session.query(question).filter(question.id==id).first()

        except:
            pass

        session.close()
        return question_result
    def genwordpdjdstr(self,tihao,id):
        single_question_set = self.gensinglequestion(id)
        self.contentstyle(str(tihao) + ':' + single_question_set.content)
    def genwordxzmxzstr(self,tihao,id):
        single_question_set = self.gensinglequestion(id)
        self.contentstyle(str(tihao) + ':' + single_question_set.content)
        if single_question_set.contentimg:
            imgpath=os.getcwd()+'\\'+single_question_set.contentimg
            self.contentimgstyle(imgpath)

        xuanxiangcontent = ''
        if len(single_question_set.choice_a) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_a
        if len(single_question_set.choice_b) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_b
        if len(single_question_set.choice_c) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_c
        if len(single_question_set.choice_d) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_d
        if len(single_question_set.choice_e) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_e
        if len(single_question_set.choice_f) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_f
        if len(single_question_set.choice_g) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_g
        if len(single_question_set.choice_h) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_h

        if not xuanxiangcontent=='':
            self.contentstyle(xuanxiangcontent)