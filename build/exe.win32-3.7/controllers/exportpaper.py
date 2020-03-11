import os

from PyQt5 import QtWidgets
from PyQt5.QtGui import QIcon
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm
from sqlalchemy.orm import sessionmaker

from controllers.utils.loginutil import CommonUtil
from model.createdb import engine
import docx
from PyQt5.QtWidgets import QWidget, QFileDialog
from views.exportpaper import  Ui_Dialog
from controllers.utils.createpaper import createpaper
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class exportpaperform(QWidget,Ui_Dialog):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowIcon(QIcon(CommonUtil.APP_ICON))
        self.initcomboBox()
        self.pushButton.clicked.connect(self.exportpapers)
        self.pushButton_2.clicked.connect(self.exportpaperformat)
    def exportpaperformat(self):
        try:
            startnum = int(self.lineEdit_2.text())-1
            endnum = int(self.lineEdit_3.text())
            self.coursename = self.comboBox_2.currentText()
            paper = createpaper(self.coursename)

            questionidformat = paper.createpaperformat(startnum,endnum)
            self.exportpaper2(questionidformat,startnum)

        except:
            QtWidgets.QMessageBox.information(self, '出卷', '请输入正确数字')

    def exportpaper2(self,questionidformat,papertihao):
        self.file = docx.Document()  # 创建内存中的word文档对象

        self.file.styles['Normal'].font.name = u'宋体'
        tihao=papertihao
        self.erjistyle('选择题')
        for i in questionidformat:
            single_question_set = self.gensinglequestion(i.id)
            if single_question_set.questionType=='xz':
                tihao += 1
                self.genwordxzmxzstr(tihao, i.id)
        self.erjistyle('多选题')
        for i in questionidformat:
            single_question_set = self.gensinglequestion(i.id)
            if single_question_set.questionType == 'mxz':
                tihao += 1
                self.genwordxzmxzstr(tihao, i.id)
        self.erjistyle('判断题')
        for i in questionidformat:
            single_question_set = self.gensinglequestion(i.id)
            if single_question_set.questionType == 'pd':
                tihao += 1
                self.genwordpdjdstr(tihao, i.id)
        self.erjistyle('填空题')
        for i in questionidformat:
            single_question_set = self.gensinglequestion(i.id)
            if single_question_set.questionType == 'jd':
                tihao += 1
                self.genwordpdjdstr(tihao, i.id)
        fileName, ok2 = QFileDialog.getSaveFileName(self,
                                                    "文件保存",
                                                    "./",
                                                    "All Files (*);;Text Files (*.docx)")

        self.filepath, self.filename = os.path.split(fileName)
        self.file.save(fileName)  # 保存才能看到结果
        filenameans= self.filepath+'/答案'+self.filename
        self.exportpaper2ans(filenameans,questionidformat,papertihao)

    def exportpaper2ans(self, filenameans,questionidformat, tihao):
        self.file = docx.Document()  # 创建内存中的word文档对象

        self.file.styles['Normal'].font.name = u'宋体'
        tihao = tihao
        xzstr=''
        for i in questionidformat:
            tihao += 1
            single_question_set = self.gensinglequestion(i.id)
            xzstr += str(tihao) + ':' + single_question_set.answer + '  '

        self.contentstyle(xzstr)

        self.file.save(filenameans)  # 保存才能看到结果

    def exportpapers(self):

        try:
             papernum=int(self.lineEdit.text())
             fileName, ok2 = QFileDialog.getSaveFileName(self,
                                                         "文件保存",
                                                         "./",
                                                         "All Files (*);;Text Files (*.docx)")
             self.filepath, self.filename = os.path.split(fileName)
             self.coursename = self.comboBox.currentText()
             for i in range(1, papernum + 1):
                 paper = createpaper(self.coursename)
                 questionidlist = paper.createpaper()
                 questionidlist.pop()
                 self.exportpaper(str(i), questionidlist)
        except:
            QtWidgets.QMessageBox.information(self, '出卷', '请输入正确数字')



    def titlestyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.bold = True
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size=Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def erjistyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    def contentstyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.first_line_indent = Cm(0.74)
    def shijianstyle(self,str):
        p = self.file.add_paragraph()
        run = p.add_run(str)
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    def genwordxzmxzstr(self,tihao,id):
        single_question_set = self.gensinglequestion(id)
        self.contentstyle(str(tihao) + ':' + single_question_set.content)
        xuanxiangcontent = ''
        if len(single_question_set.choice_a) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_a
        if len(single_question_set.choice_b) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_b
        if len(single_question_set.choice_c) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_c
        if len(single_question_set.choice_d) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_d
        if len(single_question_set.choice_e) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_e
        if len(single_question_set.choice_f) > 2:
            xuanxiangcontent += ' ' + single_question_set.choice_f
        self.contentstyle(xuanxiangcontent)
    def genwordpdjdstr(self,tihao,id):
        single_question_set = self.gensinglequestion(id)
        self.contentstyle(str(tihao) + ':' + single_question_set.content)
    def exportpaper(self,papername,questionidlist):
        self.file = docx.Document()  # 创建内存中的word文档对象

        self.file.styles['Normal'].font.name = u'宋体'
        # file.add_paragraph("窗前明月光")  # 写入若干段落
        papertitle = self.comboBox.currentText() + '卷' + papername
        if self.comboBox.currentText() is not None:



            self.titlestyle( papertitle)
            # self.erjistyle('不知道')
            # self.contentstyle('内容')
            tihao=0

            paperlistset=self.getpaperset()
            xznum=paperlistset.single_choice_num
            xzfenshu=paperlistset.single_choice_score
            pdnum=paperlistset.judgment
            pdfenshu=paperlistset.judgment_score
            mxznum=paperlistset.multiple_choice_num
            mxzfenshu=paperlistset.multiple_choice_score
            jdnum=paperlistset.jd_choice_num
            jdfenshu=paperlistset.jd_choice_score
            shijian=paperlistset.kaoshishijian
            self.shijianstyle('(考试时间：'+str(shijian)+'分钟)')
            self.erjistyle('一、选择题（每题%s分）'%xzfenshu)
            for i in range(xznum):
                tihao+=1
                id=questionidlist[i].id

                self.genwordxzmxzstr(tihao,id)

            self.erjistyle('二、判断题（每题%s分）' % pdfenshu)
            for i in range(xznum,xznum+pdnum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordpdjdstr(tihao,id)
            self.erjistyle('三、多选题（每题%s分）' % mxzfenshu)
            for i in range(xznum+pdnum,xznum+pdnum+mxznum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordxzmxzstr(tihao,id)
            self.erjistyle('四、填空题（每题%s分）' % jdfenshu)
            for i in range(xznum+pdnum+mxznum, xznum+pdnum+mxznum+jdnum):
                tihao += 1
                id = questionidlist[i].id
                self.genwordpdjdstr(tihao,id)



        filename = self.filepath + '/' +  papertitle+self.filename
        self.file.save(filename)  # 保存才能看到结果
        self.exportpaperanswer(papertitle,questionidlist)
    def exportpaperanswer(self,papertitle,questionidlist):
        self.file = docx.Document()  # 创建内存中的word文档对象

        self.file.styles['Normal'].font.name = u'宋体'
        # file.add_paragraph("窗前明月光")  # 写入若干段落



        self.titlestyle(papertitle+'答案')
        # self.erjistyle('不知道')
        # self.contentstyle('内容')
        tihao=0

        paperlistset=self.getpaperset()
        xznum=paperlistset.single_choice_num
        xzfenshu=paperlistset.single_choice_score
        pdnum=paperlistset.judgment
        pdfenshu=paperlistset.judgment_score
        mxznum=paperlistset.multiple_choice_num
        mxzfenshu=paperlistset.multiple_choice_score
        jdnum=paperlistset.jd_choice_num
        jdfenshu=paperlistset.jd_choice_score
        xzstr=''
        self.erjistyle('一、选择题（每题%s分）' % xzfenshu)
        for i in range(xznum):
            tihao+=1
            id=questionidlist[i].id
            single_question_set=self.gensinglequestion(id)
            xzstr+=str(tihao)+':'+single_question_set.answer+'  '
        self.contentstyle( xzstr)
        pdstr=''
        self.erjistyle('二、判断题（每题%s分）' % pdfenshu)
        for i in range(xznum,xznum+pdnum):
            tihao += 1
            id = questionidlist[i].id
            single_question_set = self.gensinglequestion(id)
            pdstr += str(tihao) + ':' + single_question_set.answer + '  '
        self.contentstyle( pdstr)
        mxzstr=''
        self.erjistyle('三、多选题（每题%s分）' % mxzfenshu)
        for i in range(xznum+pdnum,xznum+pdnum+mxznum):
            tihao += 1
            id = questionidlist[i].id
            single_question_set = self.gensinglequestion(id)
            mxzstr += str(tihao) + ':' + single_question_set.answer + '  '
        self.contentstyle(mxzstr)
        jdstr=''
        self.erjistyle('四、填空题（每题%s分）' % jdfenshu)
        for i in range(xznum+pdnum+mxznum, xznum+pdnum+mxznum+jdnum):
            tihao += 1
            id = questionidlist[i].id
            single_question_set = self.gensinglequestion(id)
            jdstr += str(tihao) + ':' + single_question_set.answer + '  '
        self.contentstyle( jdstr)


        filename = self.filepath + '/' + papertitle+'答案' + self.filename
        self.file.save(filename)  # 保存才能看到结果

    def getpaperset(self):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import PaperList
        question_result = None
        try:
            question_result = session.query(PaperList).filter(PaperList.course_name==self.coursename).first()

        except:
            pass

        session.close()
        return question_result
    def gensinglequestion(self,id):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import question
        question_result=None
        try:
            question_result = session.query(question).filter(question.id==id).first()

        except:
            pass

        session.close()
        return question_result



    def initcomboBox(self):
        Session = sessionmaker(bind=engine)
        # 每次执行数据库操作时，都需要创建一个session
        session = Session()
        from model.question import courselist
        try:
            result = session.query(courselist).all()
            # for single in result:
            #     # print('username:%s' % single.coursename)
            self.comboBox.addItems( i.coursename for i in result)
            self.comboBox_2.addItems( i.coursename for i in result)
        except:
            pass
        session.close()











def excel_into_model(model_name, excel_file):
    Session = sessionmaker(bind=engine)

    # 每次执行数据库操作时，都需要创建一个session
    session = Session()
    table = excel_file.sheets()[0]  # 获取第一签页
    rows = table.nrows  # 行数
    cols = table.ncols  # 列数
    colnames = table.row_values(0)  # 获取第一行一般是类别名字

    field_name = []
    # 只导入第一个sheet中的数据
    table = excel_file.sheet_by_index(0)
    nrows = table.nrows
    table_header = table.row_values(0)
    from model.question import question
    for cell in table_header:
        field_name.append(cell)
    try:
        for x in range(1, nrows):
            # 行的数据,创建对象,进行报错数据
            obj=question()
            print(len(field_name))
            for y in range(len(field_name)):
                tempfildname=field_name[y]
                cell_value=table.cell_value(x, y)
                tempstr = 'obj.%s' % field_name[y] + '=cell_value'
                exec(tempstr)
            session.add(obj)
            session.commit()
    except:
        pass
    finally:
        session.close()






